import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from colorama import Fore, Style
import time
import requests
from bs4 import BeautifulSoup


def scrape_google_scholar(topic):
    """Scrape Google Scholar for research papers"""
    try:
        print(f"{Fore.CYAN}  [*] Searching Google Scholar for research papers...{Style.RESET_ALL}")
        
        url = f"https://scholar.google.com/scholar?q={topic.replace(' ', '+')}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            papers = []
            results = soup.find_all('div', {'class': 'gs_ri'})[:10]
            
            for result in results:
                title_elem = result.find('h3', {'class': 'gs_rt'})
                snippet_elem = result.find('div', {'class': 'gs_rs'})
                
                if title_elem and snippet_elem:
                    title = re.sub(r'\[.*?\]', '', title_elem.get_text()).strip()
                    snippet = snippet_elem.get_text().strip()
                    papers.append(f"**{title}**\n{snippet}")
            
            content = '\n\n'.join(papers)
            print(f"{Fore.GREEN}  [+] Found {len(papers)} research papers{Style.RESET_ALL}")
            return content
        else:
            return ""
            
    except Exception as e:
        print(f"{Fore.YELLOW}  [!] Scholar scraping failed: {str(e)[:60]}{Style.RESET_ALL}")
        return ""


def scrape_arxiv(topic):
    """Search arXiv for academic papers"""
    try:
        print(f"{Fore.CYAN}  [*] Searching arXiv for academic papers...{Style.RESET_ALL}")
        
        # arXiv API
        url = f"http://export.arxiv.org/api/query?search_query=all:{topic.replace(' ', '+')}&start=0&max_results=10"
        
        response = requests.get(url, timeout=10)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'xml')
            entries = soup.find_all('entry')
            
            papers = []
            for entry in entries:
                title = entry.find('title')
                summary = entry.find('summary')
                
                if title and summary:
                    papers.append(f"**{title.get_text().strip()}**\n{summary.get_text().strip()}")
            
            content = '\n\n'.join(papers)
            print(f"{Fore.GREEN}  [+] Found {len(papers)} arXiv papers{Style.RESET_ALL}")
            return content
        else:
            return ""
            
    except Exception as e:
        print(f"{Fore.YELLOW}  [!] arXiv search failed: {str(e)[:60]}{Style.RESET_ALL}")
        return ""


def scrape_pubmed(topic):
    """Search PubMed for medical/biological research"""
    try:
        print(f"{Fore.CYAN}  [*] Searching PubMed for medical research...{Style.RESET_ALL}")
        
        # PubMed search
        search_url = f"https://pubmed.ncbi.nlm.nih.gov/?term={topic.replace(' ', '+')}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            articles = []
            results = soup.find_all('article', {'class': 'full-docsum'})[:10]
            
            for result in results:
                title_elem = result.find('a', {'class': 'docsum-title'})
                snippet_elem = result.find('div', {'class': 'full-view-snippet'})
                
                if title_elem:
                    title = title_elem.get_text().strip()
                    snippet = snippet_elem.get_text().strip() if snippet_elem else ""
                    articles.append(f"**{title}**\n{snippet}")
            
            content = '\n\n'.join(articles)
            print(f"{Fore.GREEN}  [+] Found {len(articles)} PubMed articles{Style.RESET_ALL}")
            return content
        else:
            return ""
            
    except Exception as e:
        print(f"{Fore.YELLOW}  [!] PubMed search failed: {str(e)[:60]}{Style.RESET_ALL}")
        return ""


def scrape_free_books(topic):
    """Search for free books on OpenLibrary and Project Gutenberg"""
    try:
        print(f"{Fore.CYAN}  [*] Searching for free books and textbooks...{Style.RESET_ALL}")
        
        books = []
        
        # OpenLibrary API
        try:
            url = f"https://openlibrary.org/search.json?q={topic.replace(' ', '+')}&limit=5"
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                docs = data.get('docs', [])
                
                for doc in docs:
                    title = doc.get('title', '')
                    author = doc.get('author_name', ['Unknown'])[0]
                    first_publish = doc.get('first_publish_year', 'N/A')
                    
                    if title:
                        books.append(f"**{title}** by {author} ({first_publish})")
        except:
            pass
        
        # Project Gutenberg search
        try:
            url = f"https://www.gutenberg.org/ebooks/search/?query={topic.replace(' ', '+')}"
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(url, headers=headers, timeout=10)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                results = soup.find_all('li', {'class': 'booklink'})[:5]
                
                for result in results:
                    title_elem = result.find('span', {'class': 'title'})
                    author_elem = result.find('span', {'class': 'subtitle'})
                    
                    if title_elem:
                        title = title_elem.get_text().strip()
                        author = author_elem.get_text().strip() if author_elem else 'Unknown'
                        books.append(f"**{title}** by {author} (Project Gutenberg)")
        except:
            pass
        
        content = '\n'.join(books)
        if books:
            print(f"{Fore.GREEN}  [+] Found {len(books)} free books{Style.RESET_ALL}")
        return content
        
    except Exception as e:
        print(f"{Fore.YELLOW}  [!] Book search failed: {str(e)[:60]}{Style.RESET_ALL}")
        return ""


def create_research_document_from_data(topic, research_data, output_path):
    """
    Create a formatted Word document from already-extracted research data
    
    Args:
        topic: Research topic
        research_data: Dictionary with research sections from Gemini
        output_path: Path to save the Word document
    """
    print(f"{Fore.CYAN}[*] Creating supplemental research document...{Style.RESET_ALL}")
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading(f'Supplemental Research: {topic}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(68, 114, 196)  # Blue
    
    # Add subtitle
    subtitle = doc.add_paragraph('Comprehensive Master Thesis-Level Research Content')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Add introduction
    intro = doc.add_paragraph(
        "This document contains comprehensive, thesis-level research content generated "
        "by AI analysis. The content includes detailed explanations, mechanisms, "
        "clinical/practical information, diagnostic procedures, treatment protocols, "
        "and current research developments (2024-2026). Use this as supplemental "
        "material for in-depth study and presentation enhancement."
    )
    intro_run = intro.runs[0]
    intro_run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # Extract sections from research data
    if research_data and research_data.get('sections'):
        for i, section in enumerate(research_data['sections'], 1):
            section_title = section.get('title', f'Section {i}')
            section_content = section.get('content', [])
            
            if not section_content:
                continue
            
            # Add section heading
            heading = doc.add_heading(f'{i}. {section_title}', 1)
            heading_run = heading.runs[0]
            heading_run.font.color.rgb = RGBColor(46, 80, 144)
            
            # Add content as paragraphs
            for bullet in section_content:
                para = doc.add_paragraph()
                para.add_run('- ').bold = True
                para.add_run(bullet)
                para_run = para.runs[1]
                para_run.font.size = Pt(11)
                
                # Add spacing between bullets
                para.paragraph_format.space_after = Pt(6)
            
            doc.add_paragraph()  # Section spacing
            
            # Page break after every 2 sections for readability
            if i % 2 == 0 and i < len(research_data['sections']):
                doc.add_page_break()
    
    # If we have raw text, add it as well
    if research_data and research_data.get('raw_text') and len(research_data['raw_text']) > 1000:
        doc.add_page_break()
        heading = doc.add_heading('Complete Research Text', 1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(46, 80, 144)
        
        # Split into paragraphs
        paragraphs = research_data['raw_text'].split('\n\n')
        for para_text in paragraphs:
            if len(para_text.strip()) > 50:
                para = doc.add_paragraph(para_text.strip())
                para_run = para.runs[0]
                para_run.font.size = Pt(11)
    
    # Add footer with generation info
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"\n\nDocument generated by AI PowerPoint Generator v2.0\n"
        f"Topic: {topic}\n"
        f"Content Type: Master Thesis-Level Research\n"
        f"Sections: {research_data.get('total_sections', 0)}\n"
        f"This is supplemental research material for academic reference."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    try:
        doc.save(output_path)
        print(f"{Fore.GREEN}[+] Research document created: {os.path.basename(output_path)}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"{Fore.RED}[!] Error saving research document: {e}{Style.RESET_ALL}")
        return False


def generate_supplemental_research(topic, research_data, output_dir):
    """
    Main function to generate supplemental research document from existing data + web sources
    
    Args:
        topic: Research topic
        research_data: Already-extracted research data from Gemini
        output_dir: Directory to save the document
    
    Returns:
        Path to the created document or None
    """
    try:
        print(f"{Fore.CYAN}[*] Gathering supplemental research from academic sources...{Style.RESET_ALL}\n")
        
        # Scrape additional sources
        scholar_content = scrape_google_scholar(topic)
        arxiv_content = scrape_arxiv(topic)
        pubmed_content = scrape_pubmed(topic)
        books_content = scrape_free_books(topic)
        
        # Add scraped content to research_data
        if not research_data:
            research_data = {'sections': [], 'raw_text': '', 'total_sections': 0}
        
        # Append additional sections
        additional_sections = []
        
        if scholar_content:
            additional_sections.append({
                'title': 'Research Papers (Google Scholar)',
                'content': scholar_content.split('\n\n')
            })
        
        if arxiv_content:
            additional_sections.append({
                'title': 'Academic Papers (arXiv)',
                'content': arxiv_content.split('\n\n')
            })
        
        if pubmed_content:
            additional_sections.append({
                'title': 'Medical Research (PubMed)',
                'content': pubmed_content.split('\n\n')
            })
        
        if books_content:
            additional_sections.append({
                'title': 'Free Books and Textbooks',
                'content': books_content.split('\n')
            })
        
        # Merge with existing sections
        if 'sections' not in research_data:
            research_data['sections'] = []
        research_data['sections'].extend(additional_sections)
        research_data['total_sections'] = len(research_data['sections'])
        
        print(f"\n{Fore.GREEN}[+] Gathered content from {len(additional_sections)} additional sources{Style.RESET_ALL}\n")
        
        # Create Word document from combined data
        safe_topic = topic.strip().replace(" ", "_").replace("\\", "").replace("/", "")
        doc_filename = os.path.join(output_dir, f'{safe_topic}_supplemental_research.docx')
        
        if create_research_document_from_data(topic, research_data, doc_filename):
            return doc_filename
        else:
            return None
            
    except Exception as e:
        print(f"{Fore.RED}[!] Error generating supplemental research: {e}{Style.RESET_ALL}")
        return None

