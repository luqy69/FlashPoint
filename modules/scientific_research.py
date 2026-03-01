"""
Scientific Research Module
Integrates free academic APIs for Super Research Mode:
- CORE: Academic papers
- ITIS: Biological taxonomy
- OpenLibrary: Book metadata
- Internet Archive: Full book text download
- LibGen: Book PDF downloads
"""

import requests
import time
import os
import sys
from colorama import Fore, Style

# Fix for Windows console encoding
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ScientificResearcher:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
    
    def search_papers(self, topic, max_results=3):
        """
        Search CORE for academic papers (NO API KEY REQUIRED)
        Returns list of papers with title, authors, year, abstract
        """
        try:
            print(f"{Fore.CYAN}[*] Searching academic papers on CORE...{Style.RESET_ALL}")
            
            # CORE API v3 (free tier, no key)
            url = "https://core.ac.uk:443/api-v2/articles/search"
            params = {
                'query': topic,
                'page': 1,
                'pageSize': max_results
            }
            
            response = self.session.get(url, params=params, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                papers = []
                
                for item in data.get('data', []):
                    papers.append({
                        'title': item.get('title', 'Unknown'),
                        'authors': ', '.join(item.get('authors', [])),
                        'year': item.get('year', 'N/A'),
                        'abstract': item.get('abstract', '')[:500],  # First 500 chars
                        'url': item.get('downloadUrl', '')
                    })
                
                print(f"{Fore.GREEN}[+] Found {len(papers)} papers{Style.RESET_ALL}")
                return papers
            else:
                print(f"{Fore.YELLOW}[!] CORE API returned status {response.status_code}{Style.RESET_ALL}")
                return []
                
        except Exception as e:
            print(f"{Fore.YELLOW}[!] CORE API error: {e}{Style.RESET_ALL}")
            return []
    
    def search_taxonomy(self, species_name):
        """
        Get biological classification from ITIS (NO API KEY REQUIRED)
        Returns taxonomy data: kingdom, phylum, class, order, family, genus, species
        """
        try:
            print(f"{Fore.CYAN}[*] Searching biological taxonomy (ITIS)...{Style.RESET_ALL}")
            
            # ITIS JSON service
            url = f"https://www.itis.gov/ITISWebService/jsonservice/searchByScientificName"
            params = {'srchKey': species_name}
            
            response = self.session.get(url, params=params, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                if data and 'scientificNames' in data:
                    tsn = data['scientificNames'][0].get('tsn')
                    
                    # Get full hierarchy
                    hierarchy_url = f"https://www.itis.gov/ITISWebService/jsonservice/getFullHierarchyFromTSN"
                    hierarchy_response = self.session.get(hierarchy_url, params={'tsn': tsn}, timeout=10)
                    
                    if hierarchy_response.status_code == 200:
                        hierarchy = hierarchy_response.json()
                        
                        taxonomy = {}
                        for item in hierarchy.get('hierarchyList', []):
                            rank = item.get('rankName', '').lower()
                            name = item.get('taxonName', '')
                            if rank in ['kingdom', 'phylum', 'class', 'order', 'family', 'genus', 'species']:
                                taxonomy[rank] = name
                        
                        print(f"{Fore.GREEN}[+] Found taxonomy data{Style.RESET_ALL}")
                        return taxonomy
            
            print(f"{Fore.YELLOW}[-] No taxonomy data found{Style.RESET_ALL}")
            return None
                
        except Exception as e:
            print(f"{Fore.YELLOW}[!] ITIS error: {e}{Style.RESET_ALL}")
            return None
    
    def search_books_metadata(self, topic, max_results=3):
        """
        Search OpenLibrary for book metadata (NO API KEY REQUIRED)
        Returns list of books with title, author, year, ISBN
        """
        try:
            print(f"{Fore.CYAN}[*] Searching books on OpenLibrary...{Style.RESET_ALL}")
            
            url = "https://openlibrary.org/search.json"
            params = {
                'q': topic,
                'limit': max_results
            }
            
            response = self.session.get(url, params=params, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                books = []
                
                for doc in data.get('docs', []):
                    # Get first ISBN if available
                    isbn = None
                    if 'isbn' in doc and doc['isbn']:
                        isbn = doc['isbn'][0]
                    
                    books.append({
                        'title': doc.get('title', 'Unknown'),
                        'author': ', '.join(doc.get('author_name', ['Unknown'])),
                        'year': doc.get('first_publish_year', 'N/A'),
                        'isbn': isbn
                    })
                
                print(f"{Fore.GREEN}[+] Found {len(books)} books{Style.RESET_ALL}")
                return books
            else:
                print(f"{Fore.YELLOW}[!] OpenLibrary returned status {response.status_code}{Style.RESET_ALL}")
                return []
                
        except Exception as e:
            print(f"{Fore.YELLOW}[!] OpenLibrary error: {e}{Style.RESET_ALL}")
            return []
    
    def search_libgen(self, title, author=None):
        """
        Search LibGen for book PDF download links
        Returns download URL if found, None otherwise
        """
        try:
            print(f"{Fore.CYAN}[*] Searching LibGen for: {title}{Style.RESET_ALL}")
            
            # Search on libgen.ac
            search_url = "https://libgen.ac/search.php"
            params = {
                'q': title,
                'res': 25
            }
            
            response = self.session.get(search_url, params=params, timeout=15)
            
            if response.status_code == 200:
                # Parse search results to find download link
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Find the first result row
                rows = soup.find_all('tr', valign='top')
                
                if rows:
                    # Get the first result's download link
                    first_row = rows[0]
                    download_links = first_row.find_all('a', href=True)
                    
                    for link in download_links:
                        href = link['href']
                        if 'download' in href.lower() or 'get' in href.lower():
                            full_url = href if href.startswith('http') else f"https://libgen.ac{href}"
                            print(f"{Fore.GREEN}[+] Found LibGen download link{Style.RESET_ALL}")
                            return full_url
                
                print(f"{Fore.YELLOW}[-] No download link found on LibGen{Style.RESET_ALL}")
                return None
            else:
                print(f"{Fore.YELLOW}[!] LibGen returned status {response.status_code}{Style.RESET_ALL}")
                return None
                
        except Exception as e:
            print(f"{Fore.YELLOW}[!] LibGen search error: {e}{Style.RESET_ALL}")
            return None
    
    def download_book_from_libgen(self, download_url, output_path):
        """
        Download book PDF from LibGen URL
        Returns True if successful, False otherwise
        """
        try:
            print(f"{Fore.CYAN}[*] Downloading from LibGen...{Style.RESET_ALL}")
            
            response = self.session.get(download_url, timeout=60, stream=True)
            
            if response.status_code == 200:
                with open(output_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                
                file_size = os.path.getsize(output_path)
                if file_size > 1000:  # At least 1KB
                    print(f"{Fore.GREEN}[+] Downloaded {file_size // 1024} KB{Style.RESET_ALL}")
                    return True
                else:
                    print(f"{Fore.YELLOW}[!] Downloaded file too small{Style.RESET_ALL}")
                    return False
            else:
                print(f"{Fore.YELLOW}[!] Download failed: status {response.status_code}{Style.RESET_ALL}")
                return False
                
        except Exception as e:
            print(f"{Fore.RED}[!] Download error: {e}{Style.RESET_ALL}")
            return False
    
    def download_book_text(self, isbn):
        """
        Download full text from Internet Archive (NO API KEY REQUIRED)
        Returns full book text or None
        """
        try:
            print(f"{Fore.CYAN}[*] Downloading book from Internet Archive...{Style.RESET_ALL}")
            
            # Search Internet Archive for book by ISBN
            search_url = "https://archive.org/advancedsearch.php"
            params = {
                'q': f'isbn:{isbn}',
                'output': 'json',
                'rows': 1
            }
            
            response = self.session.get(search_url, params=params, timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                docs = data.get('response', {}).get('docs', [])
                
                if docs:
                    identifier = docs[0].get('identifier')
                    
                    # Try multiple text formats
                    text_formats = [
                        f"https://archive.org/stream/{identifier}/{identifier}_djvu.txt",
                        f"https://archive.org/download/{identifier}/{identifier}.txt"
                    ]
                    
                    for text_url in text_formats:
                        try:
                            text_response = self.session.get(text_url, timeout=20)
                            if text_response.status_code == 200 and len(text_response.text) > 100:
                                print(f"{Fore.GREEN}[+] Downloaded book text ({len(text_response.text)} chars){Style.RESET_ALL}")
                                return text_response.text
                        except:
                            continue
            
            print(f"{Fore.YELLOW}[!] Could not download book text{Style.RESET_ALL}")
            return None
                
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Internet Archive error: {e}{Style.RESET_ALL}")
            return None
    
    def extract_book_insights(self, book_text, topic, max_chars=3000):
        """
        Extract relevant passages from book related to topic
        Uses simple keyword matching
        """
        try:
            print(f"{Fore.CYAN}[*] Extracting relevant passages...{Style.RESET_ALL}")
            
            keywords = topic.lower().split()
            relevant_sections = []
            
            # Split into paragraphs
            paragraphs = book_text.split('\n\n')
            
            for para in paragraphs:
                # Check if paragraph contains keywords
                if any(keyword in para.lower() for keyword in keywords):
                    #Clean up paragraph
                    cleaned = para.strip()
                    if len(cleaned) > 50:  # Skip very short paragraphs
                        relevant_sections.append(cleaned[:600])  # Max 600 chars per section
                        
                        if len('\n\n'.join(relevant_sections)) > max_chars:
                            break
            
            result = '\n\n'.join(relevant_sections[:5])  # Max 5 sections
            print(f"{Fore.GREEN}[+] Extracted {len(relevant_sections)} relevant passages{Style.RESET_ALL}")
            return result
            
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Extraction error: {e}{Style.RESET_ALL}")
            return ""
    
    def format_papers(self, papers):
        """Format papers for prompt injection"""
        if not papers:
            return "No papers found."
        
        formatted = []
        for i, paper in enumerate(papers, 1):
            formatted.append(f"{i}. **{paper['title']}** ({paper['year']})")
            formatted.append(f"   Authors: {paper['authors']}")
            if paper['abstract']:
                formatted.append(f"   Abstract: {paper['abstract']}")
            formatted.append("")
        
        return '\n'.join(formatted)
    
    def format_books(self, books, book_insights):
        """Format books and excerpts for prompt injection"""
        if not books and not book_insights:
            return "No books found."
        
        formatted = []
        
        # Add metadata
        for i, book in enumerate(books, 1):
            formatted.append(f"{i}. **{book['title']}** by {book['author']} ({book['year']})")
        
        formatted.append("")
        
        # Add excerpts if available
        if book_insights:
            formatted.append("**Book Excerpts:**")
            for insight in book_insights:
                formatted.append(f"\nFrom *{insight['title']}*:")
                formatted.append(insight['content'])
                formatted.append("")
        
        return '\n'.join(formatted)
    
    def format_taxonomy(self, taxonomy):
        """Format taxonomy for prompt injection"""
        if not taxonomy:
            return ""
        
        formatted = ["**Biological Classification:**"]
        for rank in ['kingdom', 'phylum', 'class', 'order', 'family', 'genus', 'species']:
            if rank in taxonomy:
                formatted.append(f"- {rank.capitalize()}: *{taxonomy[rank]}*")
        
        return '\n'.join(formatted)
    
    def save_research_papers(self, papers, output_dir):
        """
        Save research papers to text files
        """
        if not papers:
            print(f"{Fore.YELLOW}[!] No papers to save{Style.RESET_ALL}")
            return []
        
        try:
            os.makedirs(output_dir, exist_ok=True)
            saved_files = []
            
            for i, paper in enumerate(papers, 1):
                # Create filename
                safe_title = "".join(c for c in paper['title'] if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_title = safe_title[:50]
                if not safe_title:
                    safe_title = f"paper_{i}"
                
                filename = f"{safe_title}.txt"
                filepath = os.path.join(output_dir, filename)
                
                # Write paper info
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(f"Title: {paper['title']}\n")
                    f.write(f"Authors: {paper['authors']}\n")
                    f.write(f"Year: {paper['year']}\n")
                    if paper.get('url'):
                        f.write(f"URL: {paper['url']}\n")
                    f.write("\n" + "="*60 + "\n\n")
                    f.write(f"Abstract:\n{paper.get('abstract', 'No abstract available')}\n")
                
                saved_files.append(filepath)
                print(f"{Fore.GREEN}  [+] Saved paper: {filename}{Style.RESET_ALL}")
            
            print(f"{Fore.GREEN}[+] Saved {len(saved_files)} papers to: {output_dir}{Style.RESET_ALL}")
            return saved_files
            
        except Exception as e:
            print(f"{Fore.RED}[!] Error saving papers: {e}{Style.RESET_ALL}")
            return []
    
    def save_books_to_folder(self, books, topic, output_dir=None):
        """
        Download books from LibGen and Internet Archive, save as PDF or Word docs
        
        Args:
            books: List of book dictionaries with metadata
            topic: The research topic (for keyword extraction)
            output_dir: Base output directory (defaults to current directory)
        
        Returns:
            List of successfully saved book file paths
        """
        if not books:
            print(f"{Fore.YELLOW}[!] No books to save{Style.RESET_ALL}")
            return []
        
        if output_dir is None:
            output_dir = os.getcwd()
        
        try:
            os.makedirs(output_dir, exist_ok=True)
            print(f"{Fore.GREEN}[+] Books folder: {output_dir}{Style.RESET_ALL}")
        except Exception as e:
            print(f"{Fore.RED}[!] Could not create folder: {e}{Style.RESET_ALL}")
            return []
        
        saved_files = []
        
        # Check if BeautifulSoup is available for LibGen
        try:
            from bs4 import BeautifulSoup
            LIBGEN_AVAILABLE = True
        except ImportError:
            LIBGEN_AVAILABLE = False
            print(f"{Fore.YELLOW}[!] BeautifulSoup not installed - LibGen download disabled{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}   Install with: pip install beautifulsoup4{Style.RESET_ALL}")
        
        for i, book in enumerate(books, 1):
            try:
                print(f"\n{Fore.CYAN}{'='*60}{Style.RESET_ALL}")
                print(f"{Fore.CYAN}[*] Book {i}/{len(books)}: {book['title']}{Style.RESET_ALL}")
                print(f"{Fore.CYAN}{'='*60}{Style.RESET_ALL}")
                
                # Clean filename
                safe_title = "".join(c for c in book['title'] if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_title = safe_title[:50]
                if not safe_title:
                    safe_title = f"book_{i}"
                
                # Try LibGen first (PDFs)
                pdf_downloaded = False
                if LIBGEN_AVAILABLE:
                    libgen_url = self.search_libgen(book['title'], book.get('author'))
                    if libgen_url:
                        pdf_path = os.path.join(output_dir, f"{safe_title}.pdf")
                        pdf_downloaded = self.download_book_from_libgen(libgen_url, pdf_path)
                        if pdf_downloaded:
                            saved_files.append(pdf_path)
                            print(f"{Fore.GREEN}[+] Saved LibGen PDF: {safe_title}.pdf{Style.RESET_ALL}")
                            continue  # Skip to next book
                
                # If LibGen failed, try Internet Archive text
                if not pdf_downloaded and book.get('isbn'):
                    book_text = self.download_book_text(book['isbn'])
                    
                    if book_text and DOCX_AVAILABLE:
                        # Save as Word document
                        doc = Document()
                        
                        # Add title
                        title = doc.add_heading(book['title'], 0)
                        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add metadata
                        doc.add_paragraph()
                        meta_para = doc.add_paragraph()
                        meta_para.add_run(f"Author: ").bold = True
                        meta_para.add_run(book.get('author', 'Unknown'))
                        meta_para.add_run('\n')
                        meta_para.add_run(f"Year: ").bold = True
                        meta_para.add_run(str(book.get('year', 'N/A')))
                        meta_para.add_run('\n')
                        if book.get('isbn'):
                            meta_para.add_run(f"ISBN: ").bold = True
                            meta_para.add_run(book['isbn'])
                        
                        doc.add_paragraph()
                        doc.add_heading('Content', 1)
                        
                        # Extract and add excerpts
                        excerpts = self.extract_book_insights(book_text, topic, max_chars=5000)
                        if excerpts:
                            doc.add_paragraph().add_run('Relevant Excerpts:').italic = True
                            doc.add_paragraph()
                            for excerpt in excerpts.split('\n\n'):
                                if excerpt.strip():
                                    doc.add_paragraph(excerpt.strip())
                        else:
                            content_text = book_text[:5000]
                            doc.add_paragraph(content_text)
                            if len(book_text) > 5000:
                                doc.add_paragraph().add_run(f"\n[Truncated - Full book is {len(book_text)} characters]").italic = True
                        
                        # Save
                        docx_path = os.path.join(output_dir, f"{safe_title}.docx")
                        doc.save(docx_path)
                        saved_files.append(docx_path)
                        print(f"{Fore.GREEN}[+] Saved Internet Archive excerpts: {safe_title}.docx{Style.RESET_ALL}")
                        continue
                
                # If both failed, save metadata only
                if DOCX_AVAILABLE:
                    doc = Document()
                    title = doc.add_heading(book['title'], 0)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph()
                    meta_para = doc.add_paragraph()
                    meta_para.add_run(f"Author: ").bold = True
                    meta_para.add_run(book.get('author', 'Unknown'))
                    meta_para.add_run('\n')
                    meta_para.add_run(f"Year: ").bold = True
                    meta_para.add_run(str(book.get('year', 'N/A')))
                    
                    doc.add_paragraph()
                    note = doc.add_paragraph()
                    note.add_run("Note: ").bold = True
                    note.add_run("Full book text not available. This document contains only metadata.")
                    
                    docx_path = os.path.join(output_dir, f"{safe_title}_metadata.docx")
                    doc.save(docx_path)
                    saved_files.append(docx_path)
                    print(f"{Fore.YELLOW}[!] Saved metadata only: {safe_title}_metadata.docx{Style.RESET_ALL}")
                else:
                    print(f"{Fore.YELLOW}[!] Could not download book - python-docx not installed{Style.RESET_ALL}")
                
            except Exception as e:
                print(f"{Fore.RED}[!] Failed to process book: {e}{Style.RESET_ALL}")
                import traceback
                traceback.print_exc()
                continue
        
        print(f"\n{Fore.GREEN}{'='*60}{Style.RESET_ALL}")
        print(f"{Fore.GREEN}[+] Successfully saved {len(saved_files)} book(s){Style.RESET_ALL}")
        print(f"{Fore.GREEN}{'='*60}{Style.RESET_ALL}\n")
        return saved_files
